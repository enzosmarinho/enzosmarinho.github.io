from __future__ import annotations

import io
import zipfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
SOURCE = Path.home() / "Downloads" / "Curriculo_Enzo_Marinho_2026.docx"
PORTRAIT = ASSETS / "resume-portrait.png"
VISUAL_DOCX = ASSETS / "Enzo-Marinho-Curriculo-Visual.docx"
ATS_DOCX = ASSETS / "Enzo-Marinho-Curriculo-ATS.docx"

BLACK = "0B0B0B"
PAPER = "F4F5F2"
WHITE = "FFFFFF"
CORAL = "FF4D35"
CORAL_DARK = "D93723"
LIME = "D7FF45"
MUTED = "666660"
LIGHT = "E8E9E5"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Arial", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:shd"))
    if old is not None:
        p_pr.remove(old)
    p_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))


def add_hyperlink(paragraph, text, url, color=CORAL_DARK, bold=False, underline=True, size=9.0):
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(size_el)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullet_numbering(doc: Document, font="Arial") -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(ind)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid, width in zip(grid_cols, widths_dxa):
        grid.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def mark_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def configure_page(doc: Document, margin=0.68):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)


def configure_core_styles(doc: Document, visual=False):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.15 if visual else 9.35)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3.2 if visual else 3.8)
    normal.paragraph_format.line_spacing = 1.06 if visual else 1.08

    heading = styles["Heading 1"]
    heading.font.name = "Consolas"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    heading.font.size = Pt(9.4 if visual else 10.2)
    heading.font.bold = True
    heading.font.color.rgb = rgb(CORAL_DARK if visual else BLACK)
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(3.2)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True


def extract_portrait():
    if not SOURCE.exists():
        raise FileNotFoundError(f"Currículo fonte não encontrado: {SOURCE}")
    with zipfile.ZipFile(SOURCE) as archive:
        candidates = [name for name in archive.namelist() if name.startswith("word/media/") and name != "word/media/"]
        if not candidates:
            raise RuntimeError("O currículo fonte não contém imagem de perfil.")
        data = archive.read(max(candidates, key=lambda name: len(archive.read(name))))
    image = Image.open(io.BytesIO(data))
    image.save(PORTRAIT, "PNG")


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.add_run(text.upper())
    return paragraph


def add_role_header(doc, company, role, dates, accent=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1.5)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.45), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(paragraph.add_run(company), size=10.1, color=BLACK, bold=True)
    set_run_font(paragraph.add_run(f" | {role}"), size=9.1, color=MUTED, bold=True)
    set_run_font(paragraph.add_run(f"\t{dates}"), size=8.6, color=CORAL_DARK if accent else MUTED)


def add_bullet(doc, text, num_id, visual=False):
    paragraph = doc.add_paragraph()
    apply_bullet(paragraph, num_id)
    paragraph.paragraph_format.space_after = Pt(2.2 if visual else 2.7)
    paragraph.paragraph_format.line_spacing = 1.03 if visual else 1.06
    set_run_font(paragraph.add_run(text), size=8.85 if visual else 9.15, color=BLACK)
    return paragraph


def add_label_line(doc, label, text, visual=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2.2)
    set_run_font(paragraph.add_run(f"{label}: "), size=8.9 if visual else 9.2, color=CORAL_DARK if visual else BLACK, bold=True)
    set_run_font(paragraph.add_run(text), size=8.9 if visual else 9.2, color=BLACK)


def set_properties(doc, title, subject):
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Enzo Marinho"
    props.last_modified_by = "Enzo Marinho"
    props.keywords = "conteúdo, produção audiovisual, edição de vídeo, direção, IA, Remotion"
    props.comments = "Currículo profissional alinhado ao portfólio público de Enzo Marinho."


def build_visual():
    doc = Document()
    configure_page(doc, margin=0.68)  # compact_reference_guide + named portfolio_resume margin override
    configure_core_styles(doc, visual=True)
    set_properties(doc, "Enzo Marinho - Currículo Visual", "Direção, conteúdo e produção audiovisual")
    bullets = add_bullet_numbering(doc)

    masthead = doc.add_paragraph()
    masthead.paragraph_format.left_indent = Pt(10)
    masthead.paragraph_format.right_indent = Pt(10)
    masthead.paragraph_format.space_after = Pt(0)
    masthead.paragraph_format.tab_stops.add_tab_stop(Inches(1.0), WD_TAB_ALIGNMENT.LEFT)
    shade_paragraph(masthead, BLACK)
    picture = masthead.add_run()
    portrait_shape = picture.add_picture(str(PORTRAIT), height=Inches(0.78))
    portrait_shape._inline.docPr.set("descr", "Retrato de Enzo Marinho")
    set_run_font(masthead.add_run("\tENZO MARINHO"), size=26, color=WHITE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.left_indent = Pt(12)
    subtitle.paragraph_format.right_indent = Pt(12)
    subtitle.paragraph_format.space_after = Pt(0)
    subtitle.paragraph_format.line_spacing = 1.0
    shade_paragraph(subtitle, BLACK)
    set_run_font(subtitle.add_run("DIREÇÃO · CONTEÚDO · PRODUÇÃO AUDIOVISUAL"), name="Consolas", size=9.7, color=LIME, bold=True)
    set_run_font(subtitle.add_run("\nTransformo objetivos de negócio em vídeos claros, publicáveis e reaproveitáveis."), size=8.8, color=WHITE)

    contact = doc.add_paragraph()
    contact.paragraph_format.left_indent = Pt(12)
    contact.paragraph_format.right_indent = Pt(12)
    contact.paragraph_format.space_after = Pt(6)
    shade_paragraph(contact, BLACK)
    set_run_font(contact.add_run("Araçatuba/SP  |  (18) 98119-6746  |  enzosmarinho@hotmail.com  |  "), size=8.1, color=LIGHT)
    add_hyperlink(contact, "enzosmarinho.github.io", "https://enzosmarinho.github.io/", color=LIME, bold=True, underline=False, size=8.1)

    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120], indent_dxa=0)
    mark_row_as_header(table.rows[0])
    proof = [
        ("ATUAÇÃO ATUAL", "VOTI Software · CLT desde fev/2025"),
        ("SISTEMA DE CONTEÚDO", "Negócio Sem Filtro · cortes + teasers + captação long-form"),
        ("PROVA PÚBLICA", "1.733 visualizações · campanha VOTI no YouTube"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, proof):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, PAPER)
        set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), name="Consolas", size=7.1, color=CORAL_DARK, bold=True)
        set_run_font(p.add_run(f"\n{value}"), size=8.1, color=BLACK, bold=True)

    add_section_heading(doc, "Resumo profissional")
    summary = doc.add_paragraph()
    summary.paragraph_format.space_after = Pt(2.5)
    set_run_font(summary.add_run(
        "Estrategista de conteúdo e produtor audiovisual com atuação do briefing à entrega: estratégia, roteiro, direção, captação, edição e distribuição. "
        "Transformo temas técnicos e materiais longos em peças fáceis de entender, assistir e publicar. Uso IA para acelerar pesquisa, transcrição e curadoria; narrativa, acabamento e responsabilidade final permanecem humanos."
    ), size=9.05, color=BLACK)

    add_section_heading(doc, "Experiência")
    add_role_header(doc, "VOTI Software", "Estrategista de Conteúdo e Produção Audiovisual", "Fev 2025 - Atual", accent=True)
    add_bullet(doc, "Planejamento e produção para YouTube, Instagram, TikTok e LinkedIn, do roteiro à publicação.", bullets, visual=True)
    add_bullet(doc, "Direção, captação e edição de vídeos de produto, campanhas, bastidores e conteúdos educacionais.", bullets, visual=True)
    add_bullet(doc, "Tradução de temas de software e ERP em narrativas acessíveis para públicos não técnicos.", bullets, visual=True)

    add_role_header(doc, "Projetos independentes", "Produtor Audiovisual, Videomaker e Editor", "2025 - Atual")
    add_bullet(doc, "Negócio Sem Filtro: produção dos cortes e teasers publicados, mais captação em um episódio long-form, com seleção ranqueada de trechos apoiada por IA.", bullets, visual=True)
    add_bullet(doc, "Kayky Pitondo: estruturação de conteúdo long-form, correção de cor e adaptações verticais.", bullets, visual=True)
    add_bullet(doc, "Magnos Steel, 8848 Jiu-Jitsu e Lumiar Parfum (encerrado): anúncios, varejo, captação, esquetes e peças para redes sociais.", bullets, visual=True)

    add_section_heading(doc, "Trabalhos selecionados")
    cases_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(cases_table, [3120, 3120, 3120], indent_dxa=0)
    mark_row_as_header(cases_table.rows[0])
    selected_cases = [
        ("NEGÓCIO SEM FILTRO", "Cortes · teasers · captação long-form", "Sistema contínuo de conteúdo"),
        ("VOTI SOFTWARE", "Direção · captação · produto", "Campanhas, educação e bastidores"),
        ("KAYKY PITONDO", "Long-form · color · cortes", "Vídeo completo + adaptações verticais"),
    ]
    for cell, (client, role, outcome) in zip(cases_table.rows[0].cells, selected_cases):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, BLACK)
        set_cell_margins(cell, top=140, start=140, bottom=140, end=140)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(client), name="Consolas", size=7.2, color=LIME, bold=True)
        set_run_font(p.add_run(f"\n{role}"), size=8.2, color=WHITE, bold=True)
        set_run_font(p.add_run(f"\n{outcome}"), size=7.7, color=LIGHT)

    add_section_heading(doc, "Competências")
    add_label_line(doc, "Estratégia", "storytelling, copywriting, roteirização, SEO para YouTube e análise de métricas", visual=True)
    add_label_line(doc, "Produção", "direção, captação, gimbal, iluminação e áudio", visual=True)
    add_label_line(doc, "Pós-produção", "DaVinci Resolve, Adobe Premiere Pro, CapCut, color grading e design de som", visual=True)
    add_label_line(doc, "IA e automação", "Remotion, ffmpeg, transcrição, seleção de cortes e vídeo programático", visual=True)

    add_section_heading(doc, "Formação e portfólio")
    education = doc.add_paragraph()
    education.paragraph_format.space_after = Pt(1.5)
    set_run_font(education.add_run("Publicidade e Propaganda"), size=9.1, color=BLACK, bold=True)
    set_run_font(education.add_run(" | UniSALESIANO Araçatuba | Bacharelado em andamento, 8º período"), size=9.0, color=MUTED)
    portfolio = doc.add_paragraph()
    set_run_font(portfolio.add_run("Cases públicos, escopos e contato: "), size=8.9, color=BLACK, bold=True)
    add_hyperlink(portfolio, "enzosmarinho.github.io", "https://enzosmarinho.github.io/", color=CORAL_DARK, bold=True, size=8.9)

    doc.save(VISUAL_DOCX)


def build_ats():
    doc = Document()
    configure_page(doc, margin=0.72)  # compact_reference_guide + named ats_resume margin override
    configure_core_styles(doc, visual=False)
    set_properties(doc, "Enzo Marinho - Currículo ATS", "Estratégia de conteúdo, produção audiovisual e edição de vídeo")
    bullets = add_bullet_numbering(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    set_run_font(title.add_run("ENZO MARINHO"), size=23, color=BLACK, bold=True)
    headline = doc.add_paragraph()
    headline.paragraph_format.space_after = Pt(2)
    set_run_font(headline.add_run("ESTRATÉGIA DE CONTEÚDO | PRODUÇÃO AUDIOVISUAL | EDIÇÃO DE VÍDEO"), size=9.6, color=CORAL_DARK, bold=True)
    contact = doc.add_paragraph()
    contact.paragraph_format.space_after = Pt(4)
    set_run_font(contact.add_run("Araçatuba/SP | (18) 98119-6746 | enzosmarinho@hotmail.com | "), size=8.9, color=BLACK)
    add_hyperlink(contact, "Portfólio", "https://enzosmarinho.github.io/", color=BLACK, underline=True, size=8.9)
    set_run_font(contact.add_run(" | "), size=8.9, color=BLACK)
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/enzo-marinho-727200320", color=BLACK, underline=True, size=8.9)

    add_section_heading(doc, "Resumo profissional")
    p = doc.add_paragraph()
    set_run_font(p.add_run(
        "Estrategista de conteúdo e produtor audiovisual com experiência em estratégia, roteiro, direção, captação, edição e distribuição para marcas, produtos e profissionais. "
        "Atuo do briefing à entrega, transformando temas técnicos e materiais longos em conteúdo claro e publicável. Aplico IA a pesquisa, transcrição, curadoria e automação, com revisão humana de narrativa, imagem, áudio e contexto."
    ), size=9.25, color=BLACK)

    add_section_heading(doc, "Destaques profissionais")
    add_bullet(doc, "VOTI Software: atuação CLT em estratégia de conteúdo e produção audiovisual desde fevereiro de 2025.", bullets)
    add_bullet(doc, "Negócio Sem Filtro: produção dos cortes e teasers publicados, com seleção e ranking de trechos, além da captação em um episódio long-form.", bullets)
    add_bullet(doc, "Campanha sazonal da VOTI com 1.733 visualizações públicas no YouTube, verificada em julho de 2026.", bullets)

    add_section_heading(doc, "Experiência")
    add_role_header(doc, "VOTI Software", "Estrategista de Conteúdo e Produção Audiovisual", "Fev 2025 - Atual", accent=True)
    add_bullet(doc, "Planejamento e produção de conteúdo para YouTube, Instagram, TikTok e LinkedIn, do roteiro à publicação.", bullets)
    add_bullet(doc, "Direção, captação e edição de vídeos institucionais, educacionais, promocionais, de produto e de bastidores.", bullets)
    add_bullet(doc, "Tradução de temas técnicos de software e ERP em narrativas acessíveis para públicos não técnicos.", bullets)
    add_bullet(doc, "Criação de cortes verticais, campanhas e peças de produto, com acompanhamento de métricas e ajustes contínuos.", bullets)

    add_role_header(doc, "Projetos independentes", "Produtor Audiovisual, Videomaker e Editor", "2025 - Atual")
    add_bullet(doc, "Negócio Sem Filtro: cortes, teasers, curadoria de trechos apoiada por IA e captação em um episódio long-form.", bullets)
    add_bullet(doc, "Kayky Pitondo: estruturação e edição de conteúdo long-form, correção de cor e cortes verticais.", bullets)
    add_bullet(doc, "Magnos Steel, 8848 Jiu-Jitsu e Lumiar Parfum (projeto encerrado): anúncios, varejo, captação e conteúdo social.", bullets)

    add_section_heading(doc, "Competências")
    add_label_line(doc, "Estratégia e conteúdo", "storytelling, copywriting, roteirização, SEO para YouTube, análise de métricas")
    add_label_line(doc, "Produção audiovisual", "direção, captação, gimbal, iluminação, captação de áudio")
    add_label_line(doc, "Pós-produção", "DaVinci Resolve, Adobe Premiere Pro, CapCut, color grading, design de som")
    add_label_line(doc, "IA e automação", "Remotion, ffmpeg, transcrição, seleção de cortes, vídeo programático")

    add_section_heading(doc, "Formação")
    education = doc.add_paragraph()
    set_run_font(education.add_run("Publicidade e Propaganda"), size=9.35, color=BLACK, bold=True)
    set_run_font(education.add_run(" | UniSALESIANO Araçatuba | Bacharelado em andamento, 8º período"), size=9.25, color=BLACK)

    add_section_heading(doc, "Portfólio selecionado")
    links = doc.add_paragraph()
    set_run_font(links.add_run("Portfólio geral: "), size=9.1, color=BLACK, bold=True)
    add_hyperlink(links, "enzosmarinho.github.io", "https://enzosmarinho.github.io/", color=BLACK, size=9.1)
    set_run_font(links.add_run(" | VOTI Software: "), size=9.1, color=BLACK, bold=True)
    add_hyperlink(links, "youtube.com/@votigestao", "https://youtube.com/@votigestao", color=BLACK, size=9.1)
    set_run_font(links.add_run(" | Negócio Sem Filtro: "), size=9.1, color=BLACK, bold=True)
    add_hyperlink(links, "instagram.com/onegociosemfiltro", "https://instagram.com/onegociosemfiltro", color=BLACK, size=9.1)

    doc.save(ATS_DOCX)


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    extract_portrait()
    build_visual()
    build_ats()
    print(VISUAL_DOCX)
    print(ATS_DOCX)
    print(PORTRAIT)


if __name__ == "__main__":
    main()
